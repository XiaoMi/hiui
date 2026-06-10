#!/usr/bin/env python3
"""
UX Walkthrough - 本地 DOCX 生成标准入口

用途：
- 读取结构化 report.json
- 统一生成本地 .docx 报告
- 返回结构化 JSON 结果，供外层判断执行结果
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - import guard
    Document = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    Inches = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


SUPPORTED_SOURCES = {"code", "url", "screenshot"}
MAX_IMAGE_WIDTH_INCHES = 6.2
DOCX_FONT_CANDIDATES: list[tuple[str, list[str]]] = [
    ("PingFang SC", ["/System/Library/Fonts/PingFang.ttc"]),
    ("Hiragino Sans GB", ["/System/Library/Fonts/Hiragino Sans GB.ttc"]),
    ("Microsoft YaHei", ["C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/msyh.ttf"]),
    ("Noto Sans CJK SC", [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJKSC-Regular.otf",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJKSC-Regular.otf",
    ]),
    ("Source Han Sans SC", [
        "/usr/share/fonts/opentype/source-han-sans/SourceHanSansSC-Regular.otf",
        "/usr/share/fonts/opentype/adobe-source-han-sans/SourceHanSansSC-Regular.otf",
    ]),
    ("SimSun", ["C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/simsun.ttf"]),
    ("Arial Unicode MS", ["/Library/Fonts/Arial Unicode.ttf", "C:/Windows/Fonts/ARIALUNI.TTF"]),
]
DOCX_FONT_FALLBACK = "Arial"


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="根据结构化 report.json 生成 ux-walkthrough 本地 docx 报告")
    parser.add_argument("--report-json", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--json", action="store_true")
    return parser


def _fail(message: str, detail: str = "", *, reason: str = "tool_error", extra: dict | None = None) -> dict:
    payload = {
        "detail": detail[:500],
        "failure_reason": reason,
        "failure_stage": "docx_generation",
        "message": message,
        "ok": False,
        "status": "failed",
    }
    if extra:
        payload.update(extra)
    return payload


def _ok(message: str, *, output_path: str, file_name: str, source: str, issue_count: int, image_count: int) -> dict:
    return {
        "file_name": file_name,
        "image_count": image_count,
        "issue_count": issue_count,
        "message": message,
        "ok": True,
        "output_path": output_path,
        "source": source,
        "status": "generated",
    }


def _load_json(path: Path) -> dict:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ValueError(f"未找到 report.json：{path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"report.json 不是合法 JSON：{exc}") from exc
    except OSError as exc:
        raise ValueError(f"读取 report.json 失败：{exc}") from exc

    if not isinstance(data, dict):
        raise ValueError("report.json 根节点必须是对象")
    return data


def _text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _string_list(value: object) -> list[str]:
    if isinstance(value, list):
        result = []
        for item in value:
            text = _text(item)
            if text:
                result.append(text)
        return result
    text = _text(value)
    return [text] if text else []


def _normalize_source(raw: object) -> str:
    source = _text(raw).lower()
    if source not in SUPPORTED_SOURCES:
        raise ValueError("report.json 中的 source 必须是 code / url / screenshot 之一")
    return source


def _resolve_output_path(raw_output: str) -> Path:
    output = Path(raw_output).expanduser()
    if output.suffix.lower() != ".docx":
        raise ValueError("--output 必须以 .docx 结尾")
    return output.resolve()


def _select_docx_font_name() -> str:
    for font_name, candidates in DOCX_FONT_CANDIDATES:
        if any(Path(path).is_file() for path in candidates):
            return font_name
    return DOCX_FONT_FALLBACK


def _set_style_font(style, font_name: str, size: float | None = None) -> None:
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    if qn is None:
        return
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _set_run_font(run, font_name: str, size: float | None = None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if qn is None:
        return
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _resolve_image_path(raw: str, base_dir: Path) -> Path:
    path = Path(raw).expanduser()
    if not path.is_absolute():
        path = (base_dir / path).resolve()
    return path


def _normalize_image_entry(entry: object, base_dir: Path) -> dict:
    if isinstance(entry, str):
        path_text = _text(entry)
        caption = ""
    elif isinstance(entry, dict):
        path_text = _text(entry.get("path", ""))
        caption = _text(entry.get("caption", ""))
    else:
        raise ValueError("images 中的每一项必须是字符串路径或包含 path 的对象")

    if not path_text:
        raise ValueError("images 中存在空路径")

    path = _resolve_image_path(path_text, base_dir)
    if not path.is_file():
        raise ValueError(f"图片不存在：{path}")

    return {
        "caption": caption,
        "path": path,
    }


def _normalize_issue(item: object, index: int, base_dir: Path) -> dict:
    if not isinstance(item, dict):
        raise ValueError(f"issues[{index}] 必须是对象")

    title = _text(item.get("title", ""))
    if not title:
        raise ValueError(f"issues[{index}] 缺少 title")

    severity = _text(item.get("severity", "")).upper() or "P2"
    if severity not in {"P0", "P1", "P2"}:
        raise ValueError(f"issues[{index}] 的 severity 必须是 P0 / P1 / P2")

    description = _text(item.get("description", "")) or _text(item.get("problem_description", ""))
    if not description:
        raise ValueError(f"issues[{index}] 缺少 description")

    suggestion = _text(item.get("suggestion", "")) or _text(item.get("improvement", ""))
    location = _text(item.get("location", "")) or _text(item.get("position", ""))
    evidence_note = _text(item.get("evidence_note", "")) or _text(item.get("evidence", ""))
    category = _text(item.get("category", ""))
    raw_images = item.get("images", [])
    if isinstance(raw_images, (str, dict)):
        raw_images = [raw_images]
    elif not isinstance(raw_images, list):
        raise ValueError(f"issues[{index}] 的 images 必须是数组、字符串路径或包含 path 的对象")
    images = [_normalize_image_entry(entry, base_dir) for entry in raw_images]
    item_index = int(item.get("index", index))

    return {
        "category": category,
        "description": description,
        "evidence_note": evidence_note,
        "images": images,
        "index": item_index,
        "location": location,
        "severity": severity,
        "suggestion": suggestion,
        "title": title,
    }


def _normalize_report(data: dict, report_path: Path) -> dict:
    source = _normalize_source(data.get("source", ""))
    title = _text(data.get("title", "")) or f"{_text(data.get('project', 'UX Walkthrough')) or 'UX Walkthrough'} 体验走查报告"
    project = _text(data.get("project", ""))
    generated_at = _text(data.get("generated_at", ""))
    scope = _text(data.get("scope", "")) or _text(data.get("walkthrough_scope", ""))
    tech_stack = _text(data.get("tech_stack", ""))

    summary = data.get("summary", {}) if isinstance(data.get("summary"), dict) else {}
    overall = _text(summary.get("overall", "")) or _text(data.get("overall_evaluation", ""))
    key_points = _string_list(summary.get("key_points", [])) or _string_list(data.get("key_points", []))
    improvement_direction = _text(summary.get("improvement_direction", "")) or _text(data.get("improvement_direction", ""))

    issues_raw = data.get("issues", [])
    if not isinstance(issues_raw, list) or not issues_raw:
        raise ValueError("report.json 中的 issues 必须是非空数组")
    issues = [_normalize_issue(item, idx + 1, report_path.parent) for idx, item in enumerate(issues_raw)]

    appendix = data.get("appendix", {}) if isinstance(data.get("appendix"), dict) else {}
    appendix_note = _text(appendix.get("note", "")) or _text(data.get("appendix_note", ""))
    appendix_items = _string_list(appendix.get("items", []))

    return {
        "appendix_items": appendix_items,
        "appendix_note": appendix_note,
        "generated_at": generated_at,
        "improvement_direction": improvement_direction,
        "issues": issues,
        "key_points": key_points,
        "overall": overall,
        "project": project,
        "scope": scope,
        "source": source,
        "tech_stack": tech_stack,
        "title": title,
    }


def _set_base_style(document: Document) -> str:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    font_name = _select_docx_font_name()
    _set_style_font(document.styles["Normal"], font_name, 10.5)
    for style_name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12)):
        if style_name in document.styles:
            _set_style_font(document.styles[style_name], font_name, size)
    return font_name


def _add_title(document: Document, report: dict, font_name: str) -> None:
    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report["title"])
    _set_run_font(run, font_name, 18)

    meta_lines = [
        f"走查时间：{report['generated_at'] or '未提供'}",
        f"走查模式：{report['source']}",
        f"走查范围：{report['scope'] or '未提供'}",
    ]
    if report["project"]:
        meta_lines.insert(0, f"项目：{report['project']}")
    if report["tech_stack"]:
        meta_lines.append(f"技术栈：{report['tech_stack']}")

    for line in meta_lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)


def _add_summary(document: Document, report: dict) -> None:
    document.add_heading("体验问题总结", level=1)
    if report["overall"]:
        document.add_heading("整体评价", level=2)
        document.add_paragraph(report["overall"])

    if report["key_points"]:
        document.add_heading("核心问题", level=2)
        for item in report["key_points"]:
            document.add_paragraph(item, style="List Bullet")

    if report["improvement_direction"]:
        document.add_heading("改进方向", level=2)
        document.add_paragraph(report["improvement_direction"])


def _issue_heading(issue: dict) -> str:
    title = issue["title"]
    prefix = f"{issue['index']}. [{issue['severity']}]"
    if title.startswith(prefix):
        return title
    return f"{prefix}{title}" if title.startswith(" ") else f"{prefix} {title}"


def _add_issue_field(document: Document, label: str, value: str) -> None:
    if not value:
        return
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}：").bold = True
    paragraph.add_run(value)


def _add_issue_images(document: Document, issue: dict) -> int:
    image_count = 0
    for image in issue["images"]:
        document.add_picture(str(image["path"]), width=Inches(MAX_IMAGE_WIDTH_INCHES))
        image_count += 1
        if image["caption"]:
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(image["caption"]).italic = True
    return image_count


def _add_issue_detail(document: Document, issue: dict) -> int:
    document.add_heading(_issue_heading(issue), level=2)
    if issue["category"]:
        _add_issue_field(document, "问题类型", issue["category"])
    _add_issue_field(document, "位置", issue["location"])
    _add_issue_field(document, "问题描述", issue["description"])
    _add_issue_field(document, "改进建议", issue["suggestion"])
    _add_issue_field(document, "证据说明", issue["evidence_note"])
    return _add_issue_images(document, issue)


def _distribution_rows(issues: list[dict]) -> tuple[list[tuple[str, int]], list[tuple[str, int, int, int]]]:
    category_counts: dict[str, int] = {}
    severity_by_category: dict[str, dict[str, int]] = {}

    for issue in issues:
        category = issue["category"] or "未分类"
        category_counts[category] = category_counts.get(category, 0) + 1
        severity_bucket = severity_by_category.setdefault(category, {"P0": 0, "P1": 0, "P2": 0})
        severity_bucket[issue["severity"]] += 1

    top_categories = sorted(category_counts.items(), key=lambda item: (-item[1], item[0]))[:3]
    severity_rows = [
        (category, bucket["P0"], bucket["P1"], bucket["P2"])
        for category, bucket in sorted(severity_by_category.items(), key=lambda item: item[0])
    ]
    return top_categories, severity_rows


def _add_distribution(document: Document, issues: list[dict]) -> None:
    top_categories, severity_rows = _distribution_rows(issues)
    document.add_heading("问题分布", level=1)

    if top_categories:
        document.add_heading("Top 3 高频问题类型", level=2)
        table = document.add_table(rows=1, cols=3)
        header = table.rows[0].cells
        header[0].text = "排名"
        header[1].text = "问题类型"
        header[2].text = "出现次数"
        for idx, (category, count) in enumerate(top_categories, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = category
            row[2].text = str(count)

    if severity_rows:
        document.add_heading("按问题类型 × 优先级分布", level=2)
        table = document.add_table(rows=1, cols=5)
        header = table.rows[0].cells
        header[0].text = "问题类型"
        header[1].text = "P0"
        header[2].text = "P1"
        header[3].text = "P2"
        header[4].text = "合计"

        total_p0 = total_p1 = total_p2 = 0
        for category, p0, p1, p2 in severity_rows:
            row = table.add_row().cells
            row[0].text = category
            row[1].text = str(p0)
            row[2].text = str(p1)
            row[3].text = str(p2)
            row[4].text = str(p0 + p1 + p2)
            total_p0 += p0
            total_p1 += p1
            total_p2 += p2

        total_row = table.add_row().cells
        total_row[0].text = "合计"
        total_row[1].text = str(total_p0)
        total_row[2].text = str(total_p1)
        total_row[3].text = str(total_p2)
        total_row[4].text = str(total_p0 + total_p1 + total_p2)


def _add_appendix(document: Document, report: dict, total_images: int) -> None:
    document.add_heading("附录", level=1)
    document.add_paragraph(f"走查模式：{report['source']}")
    if report["project"]:
        document.add_paragraph(f"项目：{report['project']}")
    if report["scope"]:
        document.add_paragraph(f"走查范围：{report['scope']}")
    if report["appendix_note"]:
        document.add_paragraph(report["appendix_note"])
    for item in report["appendix_items"]:
        document.add_paragraph(item, style="List Bullet")
    if total_images == 0:
        document.add_paragraph("本次报告无截图证据。")


def _generate_docx(report: dict, output_path: Path) -> dict:
    if Document is None or WD_ALIGN_PARAGRAPH is None or Inches is None or Pt is None:
        return _fail(
            "DOCX generation failed: python-docx not available",
            detail=f"缺少依赖：{DOCX_IMPORT_ERROR}",
        )

    document = Document()
    font_name = _set_base_style(document)
    _add_title(document, report, font_name)
    _add_summary(document, report)
    document.add_heading("问题详情", level=1)

    total_images = 0
    for issue in report["issues"]:
        total_images += _add_issue_detail(document, issue)

    _add_distribution(document, report["issues"])
    _add_appendix(document, report, total_images)

    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
    except OSError as exc:
        return _fail("DOCX generation failed", detail=f"写入 docx 文件失败：{exc}")
    except Exception as exc:  # pragma: no cover - library specific
        return _fail("DOCX generation failed", detail=f"docx 库执行失败：{exc}")

    return _ok(
        "DOCX generated",
        file_name=output_path.name,
        image_count=total_images,
        issue_count=len(report["issues"]),
        output_path=str(output_path),
        source=report["source"],
    )


def build_result(args: argparse.Namespace) -> dict:
    output_path = _resolve_output_path(args.output)
    if output_path.exists() and not args.overwrite:
        return _fail(
            "DOCX generation failed: output exists",
            detail=f"输出文件已存在：{output_path}；若需覆盖，请显式传 --overwrite",
            reason="tool_error",
            extra={"output_path": str(output_path)},
        )

    report_path = Path(args.report_json).expanduser().resolve()
    try:
        report = _normalize_report(_load_json(report_path), report_path)
    except ValueError as exc:
        detail = str(exc)
        if detail.startswith("图片不存在："):
            return _fail("DOCX generation failed: missing image asset", detail=detail, reason="missing_asset")
        return _fail("DOCX generation failed: invalid report content", detail=detail, reason="invalid_report")

    return _generate_docx(report, output_path)


def main() -> int:
    args = _build_parser().parse_args()
    result = build_result(args)
    if args.json or not result["ok"]:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(result["output_path"])
    return 0 if result["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
